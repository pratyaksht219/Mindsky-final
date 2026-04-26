import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_document_format(doc):
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
        # Note: A4 size is typically 8.27 x 11.69 inches
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Force Times New Roman for complex scripts as well in oxml
    # style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_heading(doc, text, level=1, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    if level == 0:
        # Title
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = align
        
    # Line spacing 1.5
    p.paragraph_format.line_spacing = 1.5

def add_paragraph(doc, text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5

def main():
    doc = Document()
    set_document_format(doc)

    # --- COVER PAGE ---
    add_heading(doc, 'Seminar Report On', level=0)
    add_heading(doc, 'Mind-Sky: AI-Powered Mental Health Companion', level=0)
    add_paragraph(doc, '\n')
    
    add_paragraph(doc, 'Submitted by', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '[Name of Student] - [Registration No.]', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '\n')
    
    add_paragraph(doc, 'Bachelor of Technology', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'IN', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Computer Science and Engineering', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '\n')
    
    add_paragraph(doc, 'Under the Supervision of', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '[Name of the Faculty]', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '[Designation]', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '\n\n')
    
    add_paragraph(doc, 'LOVELY PROFESSIONAL UNIVERSITY PUNJAB', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '([Month, Year])', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    doc.add_page_break()

    # --- DECLARATION ---
    add_heading(doc, 'DECLARATION', level=0)
    decl_text = ("I hereby declare that the seminar report titled “Mind-Sky: AI-Powered Mental Health Companion” "
                 "submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology "
                 "in Computer Science and Engineering is a record of my own work carried out during the academic session _________.\n\n"
                 "I further declare that this report has not been submitted, either in part or in full, to any other institution or university "
                 "for the award of any degree or diploma.\n\n"
                 "I confirm that the content of this report is original and prepared by me. Any references used have been duly acknowledged. "
                 "I also declare that the use of Artificial Intelligence (AI) tools, if any, has been minimal and the AI-generated content in this "
                 "report is less than 10%, ensuring that the majority of the work reflects my own understanding and effort.\n\n"
                 "I take full responsibility for the authenticity and originality of the work presented in this report.")
    add_paragraph(doc, decl_text)
    add_paragraph(doc, '\n')
    add_paragraph(doc, 'Name of the Student: _______________________', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph(doc, 'Registration Number: _______________________', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph(doc, 'Course: ____________________________________', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph(doc, 'Signature of the Student: __________________', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    add_paragraph(doc, 'Date: ______________________________________', align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    doc.add_page_break()
    
    # --- Professional Profile & Repository Details ---
    add_heading(doc, 'Professional Profile & Repository Details', level=1, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'GitHub Project Repository Link: [Insert GitHub Link]')
    add_paragraph(doc, 'LinkedIn Profile Link: [Insert LinkedIn Link]')
    
    doc.add_page_break()

    # --- CHAPTER 1: Introduction ---
    add_heading(doc, 'Chapter-1: Introduction')
    
    add_heading(doc, '1.1 Title of the Seminar Topic', level=2)
    add_paragraph(doc, 'Mind-Sky: An AI-Powered Mental Health Companion App & Infrastructure.')

    add_heading(doc, '1.2 Background and Importance of the Topic', level=2)
    add_paragraph(doc, 'In today’s fast-paced world, mental health challenges such as stress, anxiety, '
                       'and depression have become increasingly prevalent. Traditional mental health care, '
                       'while effective, often faces barriers such as accessibility, cost, and social stigma. '
                       'Mind-Sky bridges this gap by leveraging artificial intelligence and modern web technologies '
                       'to provide an accessible, AI-guided mental health support system. It acts as a premium companion '
                       'application, offering real-time conversational support, structured assessments, and mood tracking, '
                       'without replacing the need for professional clinical help.')

    add_heading(doc, '1.3 Objectives of the Seminar', level=2)
    obj_text = (
        "1. To examine the architecture of a resilient mental health application built on the MERN stack.\n"
        "2. To understand the integration of a Retrieval-Augmented Generation (RAG) pipeline for clinical context.\n"
        "3. To analyze the implementation of Machine Learning-based screening via Sentence Transformers.\n"
        "4. To explore the containerized orchestration using Docker and Kubernetes.\n"
        "5. To assess real-time application observability using Prometheus and Grafana."
    )
    add_paragraph(doc, obj_text)

    add_heading(doc, '1.4 Overview of the Approach', level=2)
    add_paragraph(doc, 'The system adopts a microservices-based architecture to decouple functionalities like the API Gateway, '
                       'AI Chat Service, ML Screening Engine, and User Backend. This ensures independent scalability and fault tolerance. '
                       'The deployment is orchestrator-managed via Kubernetes, providing self-healing environments monitored by a robust observability stack.')

    doc.add_page_break()

    # --- CHAPTER 2: Literature Review ---
    add_heading(doc, 'Chapter-2: Literature Review')
    add_paragraph(doc, 'The integration of AI in mental health has seen significant advancements. Studies indicate that conversational agents can effectively mimic empathetic dialogue, assisting users in emotional regulation when clinical resources are unavailable. Retrieval-Augmented Generation (RAG) models stand at the forefront of this evolution, allowing AI systems to draw upon verified medical texts rather than hallucinating responses. '
                       'Furthermore, natural language processing models, like Sentence Transformers, have been broadly utilized for text classification tasks, demonstrating high efficacy in identifying themes of distress, anxiety, or depression in user input.')
    add_paragraph(doc, 'On the infrastructure side, the shift from monolithic architectures to microservices has mandated the use of containerization platforms like Docker. For orchestration, Kubernetes has emerged as the industry standard, capable of managing complex deployments securely. Concurrently, continuous monitoring via tools such as Prometheus (for metrics scraping) and Grafana (for visualization) is strongly recommended by DevOps literature to maintain the high availability demanded by critical healthcare applications.')

    doc.add_page_break()

    # --- CHAPTER 3: Conceptual Study / Seminar Work ---
    add_heading(doc, 'Chapter-3: Conceptual Study / Seminar Work')
    
    add_heading(doc, '3.1 Core Concepts', level=2)
    add_paragraph(doc, 'Mind-Sky is built upon the MERN stack (MongoDB, Express, React, Node.js) to deliver a seamless user interface. '
                       'It incorporates Gamification—tracking streaks, experience points (XP), and emotional scores to keep users engaged '
                       'with their digital journaling and mood tracking routines.')

    add_heading(doc, '3.2 System Architecture and Microservices', level=2)
    add_paragraph(doc, 'The application is dismantled into specialized distinct services:')
    services_text = (
        "- API Gateway: Built with Spring Boot / Express, routing traffic securely to underlying services.\n"
        "- AI Service: A FastAPI application utilizing a Chroma Vector Store and an LLM client. It processes user chat by retrieving clinically relevant documents (like PHQ-9, GAD-7 literature) via HuggingFace Embeddings before generating empathetic responses.\n"
        "- ML Screening Engine: A FastAPI-based classifier using 'sentence-transformers/all-MiniLM-L6-v2'. It calculates cosine similarities between user text and predefined prototypes for domains such as Anxiety, ADHD, Trauma, and Depression.\n"
        "- Questionnaire Service & Backend: Handles user data persistence, evaluation logic, and historical tracking in MongoDB."
    )
    add_paragraph(doc, services_text)

    add_heading(doc, '3.3 DevOps Lifecycle and Technologies', level=2)
    add_paragraph(doc, 'Mind-Sky relies heavily on modern DevOps operations.')
    devops_text = (
        "- Containerization: Each microservice (AI, Gateway, Classifier, Frontend, Backend) is packaged into isolated Docker images.\n"
        "- Orchestration: Kubernetes manifests manage replica sets, ensuring self-healing. Minikube is used for local sandboxed clustering.\n"
        "- Observability: Prometheus is configured to scrape time-series metrics from endpoints (e.g., /metrics, /actuator/prometheus). Grafana translates these metrics into real-time dashboards."
    )
    add_paragraph(doc, devops_text)

    doc.add_page_break()

    # --- CHAPTER 4: Results and Discussion ---
    add_heading(doc, 'Chapter-4: Results and Discussion')

    add_heading(doc, '4.1 Key Observations', level=2)
    add_paragraph(doc, 'The implementation of the AI service effectively bridges conversational AI with clinical guidelines. '
                       'By injecting assessment context and relevant vector store knowledge into the prompt, the LLM consistently returns structured JSON responses detailing key findings, severity explanations, and non-prescriptive recommendations without hallucinating medical diagnoses.')

    add_heading(doc, '4.2 Conceptual Analysis', level=2)
    add_paragraph(doc, 'The ML Screening engine proved highly capable in zero-shot-like classification using cosine similarity against predefined domain prototypes. This provides an immediate risk signal to the system before deep LLM analysis. '
                       'In terms of infrastructure, running multiple isolated services dramatically decreased the blast radius of service failures compared to monolithic designs. The Node.js backend handles state and MongoDB connections robustly, while specialized services process heavy AI tasks independently.')

    add_heading(doc, '4.3 Discussion of Advantages and Limitations', level=2)
    add_paragraph(doc, 'Advantages include highly scalable infrastructure, strict guardrails on AI outputs via systematic prompting, and a premium UI promoting engagement. '
                       'Limitations revolve around the inherent dangers of AI in mental healthcare; the model relies on user-provided text, which can sometimes be ambiguous. Also, current local orchestration (Minikube) requires migration to managed cloud services (like EKS or GKE) for production-grade horizontal scaling.')

    doc.add_page_break()

    # --- CHAPTER 5: Conclusion and Future Scope ---
    add_heading(doc, 'Chapter-5: Conclusion and Future Scope')

    add_heading(doc, '5.1 Summary and Conclusion', level=2)
    add_paragraph(doc, 'The Mind-Sky project successfully demonstrates building a scalable, AI-driven mental health support ecosystem. It successfully bridges sophisticated machine learning (RAG, Sentence Transformers) with robust software engineering principles (microservices, Kubernetes, Prometheus). The outcome is an application that reliably provides supportive, guided interactions securely.')

    add_heading(doc, '5.2 Future Scope', level=2)
    add_paragraph(doc, 'Future developments can include integrating more granular crisis intervention pathways effectively handed off to human professionals. From an engineering standpoint, the implementation of automated CI/CD pipelines (e.g., GitHub Actions) for seamless image building and cluster deployment is planned. Finally, the migration of the Minikube local cluster to a highly available, managed cloud infrastructure will establish Mind-Sky as a production-ready system.')

    # Save document
    doc.save('MindSky_Seminar_Report.docx')
    print("Report generated successfully: MindSky_Seminar_Report.docx")

if __name__ == "__main__":
    main()
