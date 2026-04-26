import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_document_format(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

def add_heading(doc, text, level=1):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(24)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level == 1:
        run.font.size = Pt(16)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif level == 2:
        run.font.size = Pt(14)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)

def add_list_item(doc, text):
    p = doc.add_paragraph(text, style='List Paragraph')
    p.paragraph_format.line_spacing = 1.15
    # ensure bullets by overriding style if necessary, though python-docx usually defaults list paragraph to no-bullets unless we use List Bullet
    if 'List Bullet' in doc.styles:
        p.style = doc.styles['List Bullet']
    else:
        p.style = doc.styles['List Paragraph']

def main():
    doc = Document()
    set_document_format(doc)

    add_heading(doc, 'Mindsky: Cloud-Native Architecture and DevOps Deployment', level=0)
    add_paragraph(doc, '\n')

    # 1. Abstract
    add_heading(doc, '1. Abstract')
    add_paragraph(doc, 'This report documents the architectural design, containerization, and orchestration of Mindsky, '
                       'a comprehensive microservices-based mental health platform. The system integrates real-time AI conversational '
                       'companions, ML-based distress screening, and clinical questionnaires using technologies such as Spring Boot, '
                       'FastAPI, Node.js, and React. The entire application suite is containerized utilizing Docker and deployed across '
                       'a scalable Kubernetes cluster. To ensure high availability and deep observability, Prometheus is employed for '
                       'scraping both system and custom application metrics, alongside cAdvisor for container telemetry. Grafana is '
                       'integrated to visualize this data through tailored dashboards. This project effectively demonstrates a complete, '
                       'production-ready DevOps lifecycle.')

    # 2. Introduction
    add_heading(doc, '2. Introduction')
    add_paragraph(doc, 'In the modern era of cloud computing, complex full-stack applications require robust deployment '
                       'strategies to ensure fault tolerance, scalability, and maintainability. Traditional monolithic setups '
                       'frequently suffer from tight coupling and "works-on-my-machine" dependencies, leading to inefficient release '
                       'cycles and poor resource management.')
    add_paragraph(doc, 'Mindsky tackles these challenges by adopting a microservices architecture underpinned by modern DevOps practices. '
                       'By encapsulating services within Docker containers, the application guarantees consistency across varied environments. '
                       'Kubernetes orchestrates these containers, supplying auto-healing mechanisms, intelligent routing, and load balancing.')
    add_paragraph(doc, 'Furthermore, comprehensive monitoring is critical to maintaining service reliability. Mindsky employs Prometheus '
                       'and Grafana to ingest multidimensional time-series metrics dynamically, creating an observable infrastructure '
                       'that facilitates real-time performance insights and debugging.')

    # 3. Problem Statement
    add_heading(doc, '3. Problem Statement')
    add_paragraph(doc, 'Traditional healthcare and mental wellness applications often struggle with monolithic bottlenecks, resulting in limitations such as:')
    add_list_item(doc, 'Single points of failure causing complete system outages.')
    add_list_item(doc, 'Inability to scale specific, compute-heavy components (e.g., ML inference engines) independently.')
    add_list_item(doc, 'Lack of visibility into system bottlenecks due to poor telemetric integrations.')
    add_list_item(doc, 'Inconsistent development and production environments leading to delayed deployments.')
    add_paragraph(doc, 'Mindsky aims to resolve these issues by implementing a fully containerized deployment pipeline managed by Kubernetes, ensuring separate scalability of AI and business logic, and continuous monitoring via Prometheus and Grafana.')

    # 4. Objectives
    add_heading(doc, '4. Objectives')
    add_paragraph(doc, 'The operational objectives for the DevOps lifecycle of Mindsky include:')
    add_list_item(doc, 'To modularize the application into distinct microservices (Gateway, AI, Screening, Backend, Frontend) for independent scalability.')
    add_list_item(doc, 'To containerize all components and external databases (MongoDB, Redis, PostGIS) using optimized Dockerfiles.')
    add_list_item(doc, 'To deploy the ecosystem onto a Kubernetes cluster in a dedicated namespace, utilizing structured YAML manifests for Deployments, Services, and ConfigMaps.')
    add_list_item(doc, 'To implement end-to-end monitoring using Prometheus, configuring custom scrape targets for Spring Boot, FastAPI, and Node.js applications.')
    add_list_item(doc, 'To visualize health and performance through Grafana, focusing on resource utilization and application-specific logic limits (e.g., LLM Latency, Request throughput).')

    # 5. Scope of the Project
    add_heading(doc, '5. Scope of the Project')
    add_heading(doc, '5.1 Application Scope', level=3)
    add_paragraph(doc, 'The Mindsky platform encompasses a secure Gateway routing traffic to a React frontend, an Express/MongoDB backend managing '
                       'user data, a Spring Boot questionnaire handler, a FastAPI/ChromaDB RAG-based AI service, and an ML-based text screening Engine. '
                       'The focus is strictly on continuous integration and deployment logic for these interconnected APIs.')

    add_heading(doc, '5.2 DevOps Scope', level=3)
    add_paragraph(doc, 'The DevOps boundaries include Dockerizing source code, creating multi-container compose networks for local testing, '
                       'and translating these into declarative Kubernetes objects (Pods, Deployments, NodePort/ClusterIP Services, ConfigMaps). '
                       'It strictly covers observability instrumentation via Prometheus client libraries and cAdvisor.')

    add_heading(doc, '5.3 Limitations of Scope', level=3)
    add_paragraph(doc, 'Currently, the Kubernetes deployments run locally utilizing Minikube/Docker-Desktop without persistent volume claims '
                       '(PVCs) for all databases, meaning states might not survive pod terminations. Advanced CI/CD pipelines (GitHub Actions/Jenkins) '
                       'and automated Prometheus Alertmanager rules fall out of current scope.')

    # 6. System Architecture
    add_heading(doc, '6. System Architecture')
    add_heading(doc, '6.1 Components', level=3)
    add_list_item(doc, 'Frontend (Vite/React) scaling horizontally for client connections.')
    add_list_item(doc, 'API Gateways orchestrating microservice requests.')
    add_list_item(doc, 'AI Service running Sentence Transformers and LLM inferences, requiring specific compute allocations.')
    add_list_item(doc, 'Node Backend for persistent connection to MongoDB.')
    add_list_item(doc, 'Observability Stack: cAdvisor (node telemetry), Prometheus (time-series TSDB), and Grafana (Visualizations).')
    
    add_heading(doc, '6.2 Working Flow', level=3)
    add_paragraph(doc, 'User traffic hits the Frontend or Gateway Services exposed via Kubernetes NodePorts. '
                       'From there, internal DNS resolution forwards requests to corresponding microservices (AI, Screening, Backend) communicating '
                       'over ClusterIPs. Concurrently, Prometheus actively scrapes the `/metrics` endpoints across all these pods based on service-discovery configurations.')

    # 7. Containerization with Docker
    add_heading(doc, '7. Containerization with Docker')
    add_paragraph(doc, 'Docker ensures that "Works on my machine" translates effectively to "Works in Production".')
    
    add_heading(doc, '7.1 Dockerfiles', level=3)
    add_paragraph(doc, 'Distinct base images were utilized according to application requirements. For the AI-service, `python:3.11-slim` provides a '
                       'minimal, secure environment for HuggingFace and LangChain. The Node.js backend uses `node:20-slim`, '
                       'while Spring Boot applications leverage multi-stage builds (`maven:3.9.9-eclipse-temurin` to `eclipse-temurin:21-jdk-jammy`) '
                       'to compile and package ultra-light runtime artifacts.')

    add_heading(doc, '7.2 Docker Compose', level=3)
    add_paragraph(doc, 'For rapid orchestration outside Kubernetes, a comprehensive `docker-compose.yml` ties all 10+ services '
                       'onto a shared `mental-health-network`. It dictates dependency boot orders, injects environment variables like `MONGO_URI` '
                       'and maps ports appropriately for Host-based access.')

    # 8. Kubernetes Implementation
    add_heading(doc, '8. Kubernetes Implementation')
    add_paragraph(doc, 'A shift to Kubernetes allowed for self-healing, declarative infrastructure.')
    add_list_item(doc, 'Namespaces: All resources were isolated within a dedicated `mindsky` namespace.')
    add_list_item(doc, 'Deployments: Each microservice is paired with a ReplicaSet ensuring desired state persistence. For instance, the AI deployment mounts vector stores as EmptyDirs for ephemeral processing.')
    add_list_item(doc, 'Services: NodePort is used for external-facing entry points (Gateway on 30080, Frontend on 30007), allowing browser access, while backend logic (MongoDB:27017, Redis:6379, Screening:8081) remains strictly internal behind ClusterIPs.')

    # 9. Monitoring with Prometheus
    add_heading(doc, '9. Monitoring with Prometheus')
    add_paragraph(doc, 'Prometheus was deployed within the cluster to scrape telemetry on a 5-second interval. A `ConfigMap` specifies the targets across distinct tech stacks.')
    add_list_item(doc, 'Spring Boot Apps: Scraped at `/actuator/prometheus` (Gateway, Screening, Questionnaire).')
    add_list_item(doc, 'FastAPI Services: Scraped at `/metrics` (AI, Classifier).')
    add_list_item(doc, 'Node.js Backend: Scraped using `express-prom-bundle`.')
    add_list_item(doc, 'Containers: `cAdvisor` exposes system-level RAM/CPU bottlenecks.')

    # 10. Custom Metrics
    add_heading(doc, '10. Custom Application Metrics')
    add_paragraph(doc, 'Beyond default JVM and V8 heap sizes, applications expose custom business metrics:')
    add_paragraph(doc, 'In the AI microservice, Python\'s `prometheus_client` records `LLM_LATENCY` as a Histogram to track AI generation speeds, '
                       'and Counters for `VALIDATION_FAILURES` and `LLM_FAILURES`. In the backend, dynamic route path histograms ensure insight into endpoint popularity and status code distributions.')

    # 11. Visualization with Grafana
    add_heading(doc, '11. Visualization with Grafana')
    add_paragraph(doc, 'Grafana directly queries Prometheus over internal Kubernetes routing. Pre-configured dashboards parse Prometheus PromQL queries '
                       'to surface real-time load per pod, memory leak detection, request latencies, and total request counts. This dual-layered observability '
                       'provides comprehensive systemic insight.')

    # 12. Result and Discussion
    add_heading(doc, '12. Result and Discussion')
    add_paragraph(doc, 'The ecosystem performed exceptionally under test scenarios. Kubernetes auto-reconciled manually deleted pods in seconds. '
                       'Cross-pod communication resolved flawlessly using inner-cluster DNS. Prometheus actively ingested custom AI-LLM '
                       'latencies, directly surfacing model performance into Grafana panels without imposing heavy overhead on the actual application runtime.')

    # 13. Advantages and Limitations
    add_heading(doc, '13. Advantages and Limitations')
    add_heading(doc, '13.1 Advantages', level=3)
    add_list_item(doc, 'Polyglot Scaling: Python AI APIs scale independently from Java Gateways.')
    add_list_item(doc, 'Declarative GitOps: Environments are wholly reproducible from YAML manifests.')
    add_list_item(doc, 'Transparent Processing: Metrics provide surgical views into specific microservice behaviors.')
    add_heading(doc, '13.2 Limitations', level=3)
    add_list_item(doc, 'Steeper learning curve and resource-heavy overhead for local Minikube execution.')
    add_list_item(doc, 'Data volatility without configured Persistent Volume Claims (PVCs) for core databases.')

    # 14. Future Scope
    add_heading(doc, '14. Future Scope')
    add_paragraph(doc, 'Future enhancements include migrating Minikube to managed GKE/EKS platforms, implementing TLS/HTTPS ingress controllers, '
                       'establishing comprehensive CI/CD pipelines via GitHub Actions, mapping persistent data stores dynamically, and setting threshold-based '
                       'alerting via Prometheus Alertmanager.')

    # 15. Conclusion
    add_heading(doc, '15. Conclusion')
    add_paragraph(doc, 'Mindsky proves that complex mental-health AI algorithms can be cleanly integrated into modern, production-grade '
                       'systems. By embracing Docker and Kubernetes, the project completely decouples deployment constraints from application logic. '
                       'The coupling of these techniques with the robust Promethus/Grafana observability stack yields a resilient, self-aware platform '
                       'capable of managing heavy AI compute loads reliably.')

    # 16. References
    add_heading(doc, '16. Reference')
    add_list_item(doc, 'Merkel, D. (2014). Docker: Lightweight Linux containers. Linux Journal.')
    add_list_item(doc, 'Burns, B. et al. (2016). Borg, Omega, and Kubernetes. ACM Queue.')
    add_list_item(doc, 'Prometheus Authors. Prometheus Documentation. https://prometheus.io/docs')
    add_list_item(doc, 'Grafana Labs. Grafana Documentation. https://grafana.com/docs')
    add_list_item(doc, 'LangChain Authors. LangChain Documentation. https://python.langchain.com/docs')

    doc.save('MindSky_Devops_Final_Report.docx')
    print('Successfully generated MindSky_Devops_Final_Report.docx')

if __name__ == "__main__":
    main()
